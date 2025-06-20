import fitz
import docx
from io import BytesIO
from docx import Document
from collections import Counter
import torch

from app import nlp, tokenizer, hf_model, gemini_model  # 🟢 Use preloaded models

def extract_text(file):
    if file.name.endswith(".pdf"):
        doc = fitz.open(stream=file.read(), filetype="pdf")
        return "\n".join([page.get_text() for page in doc])
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    return ""

def is_legal_document(text):
    legal_keywords = [
        "Agreement", "Contract", "This Agreement", "Terms and Conditions", "Governing Law",
        "License", "Confidential", "Arbitration", "Obligations", "Responsibilities",
        "Verification", "Authorization", "Form I-9", "U.S. Citizenship and Immigration Services"
    ]
    score = sum(1 for word in legal_keywords if word.lower() in text.lower())
    return score >= 2  # Lower this threshold if needed


def segment_clauses(text):
    doc = nlp(text)
    return [sent.text.strip() for sent in doc.sents if len(sent.text.strip()) > 20]

def classify_clause(text):
    tokens = tokenizer(text, return_tensors="pt", truncation=True, padding=True)
    with torch.no_grad():
        output = hf_model(**tokens)
    scores = torch.nn.functional.softmax(output.logits, dim=-1)[0]
    label_id = torch.argmax(scores).item()
    label = hf_model.config.id2label[label_id]
    return label, scores[label_id].item()

def detect_document_type(predictions):
    label_list = [label for label, _ in predictions]
    most_common = Counter(label_list).most_common(1)
    return most_common[0][0] if most_common else "Unknown"

def suggest_clause(text, category):
    prompt = f"The following clause is related to '{category}'. Suggest a professional rewrite or improvement:\n\n{text}"
    try:
        response = gemini_model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        return f"⚠️ Gemini Error: {str(e)}"

def generate_docx_report(clauses_data):
    doc = Document()
    doc.add_heading("AI Legal Clause Validator Report", 0)
    for idx, data in enumerate(clauses_data, 1):
        doc.add_heading(f"Clause {idx} - {data['label']} (Confidence: {data['confidence']:.2f})", level=2)
        doc.add_paragraph("Original Clause:", style='Intense Quote')
        doc.add_paragraph(data['text'])
        doc.add_paragraph("Suggested Rewrite:", style='Intense Quote')
        doc.add_paragraph(data['suggestion'])
        doc.add_paragraph("\n" + "-"*50)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
